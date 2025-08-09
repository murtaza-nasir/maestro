"""
Document Converter - Handles conversion of Word documents to Markdown

Supports:
- DOCX/DOC to Markdown conversion using python-docx
- Direct processing of Markdown files (passthrough)
- Text extraction for metadata extraction
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

class DocumentConverter:
    """Handles conversion of Word documents to Markdown format."""
    
    def __init__(self):
        """Initialize the document converter."""
        try:
            import docx
            self.docx_available = True
            logger.info("python-docx library available for Word document processing")
        except ImportError:
            self.docx_available = False
            logger.warning("python-docx library not available - Word document support will be limited")
    
    def is_word_document(self, filename: str) -> bool:
        """Check if file is a Word document."""
        return filename.lower().endswith(('.docx', '.doc'))
    
    def is_markdown_file(self, filename: str) -> bool:
        """Check if file is a Markdown file."""
        return filename.lower().endswith(('.md', '.markdown'))
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported for processing."""
        return (filename.lower().endswith('.pdf') or 
                self.is_word_document(filename) or 
                self.is_markdown_file(filename))
    
    def convert_word_to_markdown(self, word_file_path: Path) -> Optional[str]:
        """
        Convert Word document to Markdown format.
        
        Args:
            word_file_path: Path to the Word document
            
        Returns:
            Markdown content as string, or None if conversion failed
        """
        if not self.docx_available:
            logger.error("python-docx not available for Word to Markdown conversion")
            return None
            
        if not word_file_path.exists():
            logger.error(f"Word file does not exist: {word_file_path}")
            return None
            
        try:
            import docx
            
            logger.info(f"Converting {word_file_path.name} to Markdown...")
            doc = docx.Document(word_file_path)
            
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Simple markdown conversion based on style
                style_name = paragraph.style.name.lower()
                
                if 'heading 1' in style_name or 'title' in style_name:
                    markdown_content.append(f"# {text}")
                elif 'heading 2' in style_name:
                    markdown_content.append(f"## {text}")
                elif 'heading 3' in style_name:
                    markdown_content.append(f"### {text}")
                elif 'heading 4' in style_name:
                    markdown_content.append(f"#### {text}")
                elif 'heading 5' in style_name:
                    markdown_content.append(f"##### {text}")
                elif 'heading 6' in style_name:
                    markdown_content.append(f"###### {text}")
                else:
                    # Regular paragraph
                    markdown_content.append(text)
                
                # Add blank line after each paragraph
                markdown_content.append("")
            
            # Handle tables
            for table in doc.tables:
                markdown_content.append("")  # Blank line before table
                
                # Table headers
                if table.rows:
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    markdown_content.append("| " + " | ".join(headers) + " |")
                    markdown_content.append("| " + " | ".join(['---'] * len(headers)) + " |")
                    
                    # Table data rows
                    for row in table.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_content.append("| " + " | ".join(cells) + " |")
                
                markdown_content.append("")  # Blank line after table
            
            result = "\n".join(markdown_content).strip()
            
            if not result:
                logger.warning(f"No content extracted from Word document: {word_file_path.name}")
                return ""
            
            logger.info(f"Successfully converted {word_file_path.name} to Markdown ({len(result)} characters)")
            return result
            
        except Exception as e:
            logger.error(f"Error converting Word document to Markdown: {e}")
            return None
    
    def read_markdown_file(self, markdown_file_path: Path) -> Optional[str]:
        """
        Read Markdown file content directly.
        
        Args:
            markdown_file_path: Path to the Markdown file
            
        Returns:
            Markdown content as string, or None if reading failed
        """
        try:
            if not markdown_file_path.exists():
                logger.error(f"Markdown file does not exist: {markdown_file_path}")
                return None
                
            with open(markdown_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            if not content.strip():
                logger.warning(f"Markdown file is empty: {markdown_file_path.name}")
                return ""
                
            logger.info(f"Successfully read Markdown file: {markdown_file_path.name} ({len(content)} characters)")
            return content
            
        except Exception as e:
            logger.error(f"Error reading Markdown file {markdown_file_path}: {e}")
            return None
    
    def extract_initial_text_for_metadata(self, file_path: Path) -> str:
        """
        Extract initial text from various document formats for metadata extraction.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Initial text content for metadata extraction
        """
        filename = file_path.name
        
        try:
            # For Word documents, extract text using python-docx
            if self.is_word_document(filename):
                if not self.docx_available:
                    logger.warning("python-docx not available for text extraction")
                    return ""
                    
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    # Extract first few paragraphs for metadata
                    paragraphs = []
                    for paragraph in doc.paragraphs[:15]:  # First 15 paragraphs
                        text = paragraph.text.strip()
                        if text:
                            paragraphs.append(text)
                            
                        # Stop if we have enough text for metadata extraction
                        if len('\n\n'.join(paragraphs)) > 2000:
                            break
                    
                    initial_text = '\n\n'.join(paragraphs)
                    
                    if initial_text:
                        logger.info(f"Extracted initial text from Word document: {filename} ({len(initial_text)} characters)")
                        return initial_text
                    else:
                        logger.warning(f"No text found in Word document: {filename}")
                        return ""
                        
                except Exception as e:
                    logger.warning(f"Failed to extract text from Word document {filename}: {e}")
                    return ""
            
            # For Markdown files, read directly
            elif self.is_markdown_file(filename):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                # Take first 2000 characters for metadata extraction
                initial_text = content[:2000] if len(content) > 2000 else content
                logger.info(f"Extracted initial text from Markdown file: {filename} ({len(initial_text)} characters)")
                return initial_text
            
            # For PDF files, return empty (existing PDF processing handles this)
            elif filename.lower().endswith('.pdf'):
                return ""  # Let existing PDF processing handle this
                
            else:
                logger.error(f"Unsupported file format for text extraction: {filename}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting initial text from {filename}: {e}")
            return ""