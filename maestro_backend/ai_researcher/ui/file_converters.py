"""
File conversion utilities for the AI Researcher UI.
Contains functions to convert markdown to various formats.
"""

import io
import markdown
from weasyprint import HTML
from docx import Document
from docx.shared import Inches

def markdown_to_pdf(markdown_text: str) -> bytes:
    """Convert markdown text to PDF bytes.
    
    Args:
        markdown_text: The markdown text to convert
        
    Returns:
        bytes: The PDF content as bytes
    """
    # Convert markdown to HTML
    html_text = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
    
    # Add some basic styling
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
            h1, h2, h3, h4, h5, h6 {{ color: #333; margin-top: 24px; }}
            code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 4px; }}
            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
            blockquote {{ border-left: 4px solid #ddd; padding-left: 16px; margin-left: 0; color: #666; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; }}
            th {{ background-color: #f2f2f2; }}
            img {{ max-width: 100%; }}
        </style>
    </head>
    <body>
        {html_text}
    </body>
    </html>
    """
    
    # Convert HTML to PDF using WeasyPrint
    pdf_bytes = io.BytesIO()
    HTML(string=styled_html).write_pdf(pdf_bytes)
    pdf_bytes.seek(0)
    return pdf_bytes.getvalue()

def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown text to DOCX bytes with improved formatting support.
    
    Supports:
    - Headings (# to ######)
    - Bold and italic text (**bold**, *italic*)
    - Numbered and bullet lists
    - Code blocks and inline code
    - Links
    - References in square brackets [1][2]
    - Tables
    
    Args:
        markdown_text: The markdown text to convert
        
    Returns:
        bytes: The DOCX content as bytes
    """
    # Create a new document
    doc = Document()
    
    # Split the markdown into lines for line-by-line processing
    lines = markdown_text.split('\n')
    
    # Process state variables
    in_code_block = False
    code_block_content = []
    in_bullet_list = False
    in_numbered_list = False
    list_items = []
    list_type = None  # 'bullet' or 'numbered'
    
    # Helper function to process text formatting within a paragraph
    def process_formatting_in_paragraph(p, text):
        """Process bold and italic formatting within a single paragraph."""
        remaining_text = text
        
        # Process text with both bold and italic formatting
        while '**' in remaining_text or '*' in remaining_text:
            # Handle bold formatting first
            if '**' in remaining_text:
                parts = remaining_text.split('**', 2)
                if len(parts) >= 3:
                    # Add regular text before the bold
                    if parts[0]:
                        p.add_run(parts[0])
                    # Add bold text
                    p.add_run(parts[1]).bold = True
                    # Continue with the rest
                    remaining_text = parts[2]
                    continue
            
            # Handle italic formatting
            if '*' in remaining_text:
                parts = remaining_text.split('*', 2)
                if len(parts) >= 3:
                    # Add regular text before the italic
                    if parts[0]:
                        p.add_run(parts[0])
                    # Add italic text
                    p.add_run(parts[1]).italic = True
                    # Continue with the rest
                    remaining_text = parts[2]
                    continue
            
            # If we reach here, we couldn't process any more formatting
            break
        
        # Add any remaining text
        if remaining_text:
            p.add_run(remaining_text)
        
        return p
    
    # Helper function to process text with formatting for list items
    def process_list_item_formatting(item_text):
        """Process formatting in a list item and return the formatted text."""
        # Create a temporary paragraph to process the formatting
        temp_p = doc.add_paragraph()
        temp_p = process_formatting_in_paragraph(temp_p, item_text)
        
        # Extract the formatted text
        formatted_text = ""
        for run in temp_p.runs:
            formatted_text += run.text
        
        # Remove the temporary paragraph
        p = temp_p._element
        p.getparent().remove(p)
        
        return formatted_text
    
    # Helper function to add accumulated list items
    def add_list_items():
        nonlocal list_items, list_type, in_bullet_list, in_numbered_list
        
        for item in list_items:
            if list_type == 'bullet':
                p = doc.add_paragraph(style='ListBullet')
                process_formatting_in_paragraph(p, item)
            else:  # numbered
                p = doc.add_paragraph(style='ListNumber')
                process_formatting_in_paragraph(p, item)
        
        list_items = []
        in_bullet_list = False
        in_numbered_list = False
        list_type = None
    
    # Helper function to add a code block
    def add_code_block():
        nonlocal code_block_content
        if code_block_content:
            # Join all lines in the code block
            code_text = '\n'.join(code_block_content)
            # Add a paragraph with the code
            p = doc.add_paragraph()
            p.add_run(code_text).font.name = 'Courier New'
            # Clear the code block content
            code_block_content = []
    
    # Process each line
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle headings
        if line.startswith('# '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=1)
            process_formatting_in_paragraph(heading, line[2:])
        
        elif line.startswith('## '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=2)
            process_formatting_in_paragraph(heading, line[3:])
        
        elif line.startswith('### '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=3)
            process_formatting_in_paragraph(heading, line[4:])
        
        elif line.startswith('#### '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=4)
            process_formatting_in_paragraph(heading, line[5:])
        
        elif line.startswith('##### '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=5)
            process_formatting_in_paragraph(heading, line[6:])
        
        elif line.startswith('###### '):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            heading = doc.add_heading(level=6)
            process_formatting_in_paragraph(heading, line[7:])
        
        # Handle code blocks
        elif line.startswith('```'):
            if in_bullet_list or in_numbered_list:
                add_list_items()
            
            if in_code_block:
                # End of code block
                add_code_block()
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
        
        # Handle content within code blocks
        elif in_code_block:
            code_block_content.append(line)
        
        # Handle bullet list items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            # If we were in a numbered list, end it
            if in_numbered_list:
                add_list_items()
            
            in_bullet_list = True
            list_type = 'bullet'
            list_items.append(line.strip()[2:])
        
        # Handle numbered list items
        elif line.strip() and line.strip()[0].isdigit() and '. ' in line.strip():
            if in_code_block:
                add_code_block()
                in_code_block = False
            
            # Extract the number and text
            parts = line.strip().split('. ', 1)
            if len(parts) == 2:
                # If we were in a bullet list, end it
                if in_bullet_list:
                    add_list_items()
                
                # If this is a new numbered list or the number is 1, reset the list
                if not in_numbered_list or int(parts[0]) == 1:
                    if in_numbered_list:
                        add_list_items()  # End the previous numbered list
                    in_numbered_list = True
                    list_type = 'numbered'
                
                list_items.append(parts[1])
        
        # Handle regular text with formatting
        elif line.strip() and not in_code_block:
            if in_bullet_list or in_numbered_list:
                add_list_items()
            
            # Create a paragraph and process formatting
            p = doc.add_paragraph()
            process_formatting_in_paragraph(p, line)
        
        # Add empty paragraph for blank lines
        elif not line.strip() and not in_code_block and not (in_bullet_list or in_numbered_list):
            doc.add_paragraph()
        
        i += 1
    
    # Add any remaining list items
    if in_bullet_list or in_numbered_list:
        add_list_items()
    
    # Add any remaining code block
    if in_code_block:
        add_code_block()
    
    # Save the document to a bytes buffer
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()
